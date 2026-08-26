from __future__ import annotations

import copy
import hashlib
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "论文写作" / "V5.0.docx"
OUTPUT = ROOT / "论文写作" / "V5.0_续写版.docx"
FIGURE_DIR = ROOT / "results" / "figures"


def set_run_font(run, size=12, bold=None, italic=None, color=None):
    run.font.name = "Times New Roman"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def set_cell_margins(cell, top=90, start=100, bottom=90, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_width(cell, width_cm):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(Cm(width_cm).emu / 635)))
    tc_w.set(qn("w:type"), "dxa")
    cell.width = Cm(width_cm)


def set_table_width(table, widths_cm):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    total_dxa = sum(int(Cm(w).emu / 635) for w in widths_cm)
    tbl_w.set(qn("w:w"), str(total_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_cm:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(Cm(width).emu / 635)))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_cm[idx])


def set_table_borders(table, color="808080", size=4):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_no_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def element_digest(elements):
    hasher = hashlib.sha256()
    for element in elements:
        hasher.update(element.xml.encode("utf-8"))
    return hasher.hexdigest()


def add_before_reference(doc, reference_element, element):
    reference_element.addprevious(element)


def add_heading(doc, reference_element, text, level):
    style = {2: "Heading 2", 3: "Heading 3", 4: "Heading 4"}[level]
    paragraph = doc.add_paragraph(style=style)
    paragraph.add_run(text)
    paragraph.paragraph_format.keep_with_next = True
    add_before_reference(doc, reference_element, paragraph._p)
    return paragraph


def add_body(doc, reference_element, text, segments=None):
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Cm(0.85)
    fmt.line_spacing = 1.5
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    if segments:
        for segment_text, bold in segments:
            run = paragraph.add_run(segment_text)
            set_run_font(run, 12, bold=bold)
    else:
        run = paragraph.add_run(text)
        set_run_font(run, 12)
    add_before_reference(doc, reference_element, paragraph._p)
    return paragraph


def add_note(doc, reference_element, text):
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Cm(0)
    fmt.line_spacing = 1.15
    fmt.space_before = Pt(2)
    fmt.space_after = Pt(4)
    run = paragraph.add_run(text)
    set_run_font(run, 9)
    add_before_reference(doc, reference_element, paragraph._p)
    return paragraph


def add_table_title(doc, reference_element, text):
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Cm(0)
    fmt.line_spacing = 1.0
    fmt.space_before = Pt(6)
    fmt.space_after = Pt(3)
    fmt.keep_with_next = True
    run = paragraph.add_run(text)
    set_run_font(run, 10.5, bold=True)
    add_before_reference(doc, reference_element, paragraph._p)
    return paragraph


def add_table(doc, reference_element, headers, rows, widths, font_size=8.5, left_columns=None):
    left_columns = set(left_columns or [])
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Normal Table"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    set_table_width(table, widths)
    set_repeat_table_header(table.rows[0])
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        shade_cell(cell, "D9EAF7")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        run = paragraph.add_run(str(header))
        set_run_font(run, font_size, bold=True)
    for row_values in rows:
        row = table.add_row()
        set_no_row_split(row)
        for idx, value in enumerate(row_values):
            cell = row.cells[idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx in left_columns else WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            run = paragraph.add_run(str(value))
            set_run_font(run, font_size)
    set_table_width(table, widths)
    add_before_reference(doc, reference_element, table._tbl)
    return table


def add_figure(doc, reference_element, filename, caption, width_cm=15.2):
    image_path = FIGURE_DIR / filename
    if not image_path.exists():
        raise FileNotFoundError(image_path)
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Cm(0)
    fmt.space_before = Pt(6)
    fmt.space_after = Pt(2)
    fmt.keep_with_next = True
    run = paragraph.add_run()
    inline_shape = run.add_picture(str(image_path), width=Cm(width_cm))
    inline_shape._inline.docPr.set("descr", caption)
    inline_shape._inline.docPr.set("title", caption)
    add_before_reference(doc, reference_element, paragraph._p)

    cap = doc.add_paragraph(style="Normal")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_fmt = cap.paragraph_format
    cap_fmt.first_line_indent = Cm(0)
    cap_fmt.line_spacing = 1.0
    cap_fmt.space_before = Pt(0)
    cap_fmt.space_after = Pt(6)
    cap_run = cap.add_run(caption)
    set_run_font(cap_run, 9)
    add_before_reference(doc, reference_element, cap._p)
    return paragraph, cap


def add_page_break(doc, reference_element):
    paragraph = doc.add_paragraph(style="Normal")
    paragraph.add_run().add_break()
    add_before_reference(doc, reference_element, paragraph._p)
    return paragraph


def append_reference(doc, text):
    paragraph = doc.add_paragraph(style="参考文献")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt = paragraph.paragraph_format
    fmt.left_indent = Cm(0.74)
    fmt.first_line_indent = Cm(-0.74)
    fmt.line_spacing = 1.0
    run = paragraph.add_run(text)
    set_run_font(run, 10.5)
    return paragraph


def build():
    doc = Document(SOURCE)
    reference_heading = next(p for p in doc.paragraphs if p.text.strip() == "参考文献")
    reference_element = reference_heading._p
    chapter5 = next(p for p in doc.paragraphs if p.text.strip() == "5 实验")

    body_children = list(doc._element.body)
    chapter5_index = body_children.index(chapter5._p)
    unchanged_elements = body_children[: chapter5_index + 1]
    before_digest = element_digest(unchanged_elements)

    if chapter5._p.getnext() is not None:
        following = chapter5._p.getnext()
        if following.tag == qn("w:p") and not "".join(following.itertext()).strip():
            following.getparent().remove(following)

    add_body(
        doc,
        reference_element,
        "本章围绕第4章提出的传播动能辨识、破圈临界预警和时滞闭环控制三个模块开展实验验证。实验首先说明数据来源、对比方法、评价指标和参数设置，随后按照“模块—实验”一一对应的方式分析传播规模预测、传播动能因子、破圈预警与控制策略的有效性，并通过多随机种子、时间切分和典型案例检验结果的稳定性与适用边界。",
    )

    add_heading(doc, reference_element, "5.1 实验设置", 3)
    add_heading(doc, reference_element, "5.1.1 实验数据", 4)
    add_body(
        doc,
        reference_element,
        "为确保实验同时覆盖真实时间演化、跨平台传播结构和中文大规模级联场景，本文选取PHEME、Twitter15、Twitter16和Weibo四个公开数据集。PHEME包含源帖、回复文本、真实时间戳、会话结构和用户属性，适合验证多模态传播规模预测、破圈预警和时滞控制；Twitter15与Twitter16提供源帖文本、传播树及相对延迟，可用于检验跨事件、跨平台的泛化能力；Weibo包含规模较大的传播结构，但当前版本缺少可用的源文本、真实时间戳和用户画像，因此仅作为结构与代理时间条件下的辅助实验，不纳入主要结论。",
    )
    add_table_title(doc, reference_element, "表1 实验数据集统计")
    add_table(
        doc,
        reference_element,
        ["数据集", "样本数", "用户数", "传播边数", "平均级联节点数", "主要用途"],
        [
            ["PHEME", "6425", "50593", "95074", "16.28", "真实时间、多模态主实验"],
            ["Twitter15", "1490", "48417", "607289", "37.63", "跨平台泛化实验"],
            ["Twitter16", "818", "23913", "352987", "33.85", "跨平台泛化实验"],
            ["Weibo", "4659", "未提供", "3677338", "790.30", "结构/代理时间辅助实验"],
        ],
        [2.1, 1.8, 2.0, 2.3, 2.5, 4.2],
        font_size=8.5,
        left_columns={5},
    )
    add_note(
        doc,
        reference_element,
        "注：PHEME中有1个标签不明确样本在监督划分时被剔除；Twitter15/16的用户画像与转发正文不完整；Weibo的用户数、真实时间戳和源文本在当前预处理版本中不可用。",
    )
    add_body(
        doc,
        reference_element,
        "所有数据均被转换为统一的样本、事件、传播边、动态图快照和社区标识格式。传播规模预测仅使用观察窗口内可获得的源帖文本、传播图、时间序列与用户统计量，最终级联规模只作为预测目标，以避免未来信息泄漏。破圈标签根据跨社区边数量、跨社区边比例和社区覆盖数的同步增长构造；当某一时间窗口首次满足预设破圈条件时，将其记为breakout time。该标签具有可解释性，但仍属于启发式标注，因此后续结果反映的是对当前破圈定义的预测能力。",
    )

    add_heading(doc, reference_element, "5.1.2 基线方法", 4)
    add_body(
        doc,
        reference_element,
        "",
        segments=[
            ("DSHCL-adapted[18]：", True),
            ("受对比超图学习方法启发，以扩散状态和交互状态构造高阶传播表示，并通过对齐后的表示连接相同预测头，用于检验高阶交互表征能否替代本文的异构动能因子。由于原始用户关系与完整高阶交互不可得，本文采用级联社区和时间窗口共现作为代理，因此将其明确标记为改编版本。", False),
        ],
    )
    add_body(
        doc,
        reference_element,
        "",
        segments=[
            ("MIDPMS-adapted[19]：", True),
            ("受多尺度扩散预测网络启发，利用宏观级联统计、微观路径和多尺度时间窗口预测最终传播规模，用于检验多尺度扩散预测在不引入完整异构闭环机制时的性能。其微观路径模块根据当前可用的传播树和窗口特征进行适配。", False),
        ],
    )
    add_body(
        doc,
        reference_element,
        "",
        segments=[
            ("SEIZ/CD-SEIZ[20]：", True),
            ("采用易感、暴露、感染/采纳和怀疑/停滞状态刻画信息扩散动力学，并在验证集上选择SIR、SEIR或SEIZ分支。该基线用于检验经典宏观动力学在短期级联规模预测和控制仿真中的解释能力。", False),
        ],
    )
    add_body(
        doc,
        reference_element,
        "",
        segments=[
            ("ED-ID-adapted[21]：", True),
            ("受社交网络信息扩散演化模型启发，以内部社区协调、外部风险激励、扩散压力和干预成本构造演化博弈控制策略，用于比较现有扩散演化思想与本文风险驱动闭环控制的差异。本文同时给出其原生预算版本和严格同预算版本，避免高成本策略因投入更多资源而获得不公平优势。", False),
        ],
    )
    add_body(
        doc,
        reference_element,
        "除上述同领域基线外，实验还设置静态图、动态趋势、仅动态图、仅社区特征、固定干预、影响力阻断、随机同预算和固定同预算等基础对照，并通过删除低频能量、跨社区拓扑、博弈响应和事件触发机制构造消融方案。所有可比较方法使用相同的数据划分、观察窗口和测试样本；模型分支与阈值仅在验证集上选择。",
    )

    add_heading(doc, reference_element, "5.1.3 评价指标", 4)
    add_body(
        doc,
        reference_element,
        "本文从预测精度、预警能力和控制效果三个层面进行评价。传播规模预测采用平均绝对误差（MAE）、均方根误差（RMSE）、平均绝对百分比误差（MAPE）和决定系数（R²）；破圈预警采用曲线下面积（AUC）、F1值、Recall@10%和平均提前时间；闭环控制采用传播抑制率、干预成本、收益成本比和触发率。其中，MAE、RMSE、MAPE和成本越小越好，R²、AUC、F1、Recall@10%、提前时间、抑制率和收益成本比越大越好。",
    )

    add_heading(doc, reference_element, "5.1.4 参数设置", 4)
    add_body(
        doc,
        reference_element,
        "主实验采用训练集、验证集和测试集7∶1∶2的固定分层划分，划分随机种子为42；稳定性实验使用7、21、42、84和2024五个模型随机种子。在此基础上，额外采用按时间排序的划分进行无泄漏压力测试。传播规模预测的主观察窗口设为180 min，并以60 min和360 min窗口检验早期信息不足与观察信息增加时的变化。",
    )
    add_table_title(doc, reference_element, "表2 主要实验参数")
    add_table(
        doc,
        reference_element,
        ["实验模块", "参数", "取值", "设置说明"],
        [
            ["通用", "数据划分", "7∶1∶2", "分层主实验；另设时间切分"],
            ["通用", "随机种子", "7/21/42/84/2024", "报告均值±标准差"],
            ["C1/V1", "观察窗口", "180 min", "60/360 min用于敏感性分析"],
            ["C1/V1", "隐藏维度/图层数", "64/2", "Dropout=0.3"],
            ["C1/V1", "学习率/权重衰减", "0.001/0.0001", "AdamW优化"],
            ["C1/V2", "潜变量维度K", "4", "按验证集MAPE选择"],
            ["C1/V2", "KL/重构权重", "0.1/0.1", "KL前5轮预热"],
            ["C2", "窗口长度/最大窗口数", "60 min/12", "按窗口滚动计算风险"],
            ["C3", "控制时滞/效果系数", "1窗口/0.85", "模拟审核与策略生效延迟"],
            ["C3", "脉冲强度范围", "0.10～0.65", "风险驱动自适应裁剪"],
        ],
        [2.1, 3.4, 3.4, 6.1],
        font_size=8.2,
        left_columns={1, 3},
    )
    add_note(
        doc,
        reference_element,
        "注：V1/V2模型批量大小为64，最大训练轮数为30～40，并采用验证集早停；C2当前实现以窗口级时序、频谱与社区特征训练可复现分类器，C3为基于C2风险分数和观测级联快照的闭环仿真。",
    )

    add_heading(doc, reference_element, "5.2 实验分析", 3)
    add_body(
        doc,
        reference_element,
        "本节按照第4章三个模块的顺序组织实验。首先检验多模态传播规模预测与传播动能因子分解，其次分析低频能量和跨社区拓扑对破圈预警的贡献，最后比较风险驱动闭环控制与固定、随机及演化博弈策略的抑制效果和成本。",
    )

    add_heading(doc, reference_element, "5.2.1 传播动能辨识有效性分析", 4)
    add_body(
        doc,
        reference_element,
        "传播动能辨识模块需要从有限观察窗口内的文本、拓扑、时序和用户信息中预测最终级联规模，并将多源表征压缩为可解释的低维因子。为验证该模块的有效性，本文首先在统一的180 min观察窗口和相同测试样本上，将HeteroRumorDyn-V1与MIDPMS-adapted、DSHCL-adapted和SEIZ/CD-SEIZ进行比较，结果如表3所示。",
    )
    add_table_title(doc, reference_element, "表3 180 min观察窗口下传播规模预测结果")
    add_table(
        doc,
        reference_element,
        ["数据集", "HeteroRumorDyn-V1", "MIDPMS-adapted", "DSHCL-adapted", "SEIZ/CD-SEIZ"],
        [
            ["PHEME", "0.1247±0.0019", "0.2222±0.0003", "0.2267±0.0040", "0.1234"],
            ["Twitter15", "0.1785±0.0019", "0.2632", "0.2526±0.0089", "0.2480"],
            ["Twitter16", "0.2010±0.0072", "0.2947", "0.2929±0.0025", "0.2635"],
        ],
        [2.2, 3.3, 3.2, 3.2, 3.1],
        font_size=8.5,
    )
    add_note(
        doc,
        reference_element,
        "注：表中数值为测试集MAPE，越小越好；带“±”的结果为五个随机种子的均值±标准差，SEIZ/CD-SEIZ为验证集选择后的单次确定性结果。",
    )
    add_body(
        doc,
        reference_element,
        "由表3可知，HeteroRumorDyn-V1在Twitter15和Twitter16上取得最低MAPE，分别为0.1785和0.2010；相对于对应数据集表现最好的同领域基线，误差分别降低约28.0%和23.7%。在PHEME上，SEIZ/CD-SEIZ的MAPE为0.1234，略低于本文模型的0.1247，但二者差距较小，配对自助法置信区间覆盖零。这说明PHEME中较短、较规则的会话级联能够被经典动力学较好拟合，而在Twitter15/16更长且结构差异更明显的传播过程中，多模态表示具有更稳定的优势。",
    )
    add_body(
        doc,
        reference_element,
        "时间切分压力测试进一步表明，HeteroRumorDyn-V1在PHEME、Twitter15和Twitter16上的MAPE分别为0.1078、0.2010和0.2581，均低于同一设置下的最优基线。该结果说明模型优势并非完全来自随机划分中的事件相似性，而是在时间分布发生变化时仍能保留一定的传播规模预测能力。",
    )
    add_figure(
        doc,
        reference_element,
        "fig5_v1_window_and_ablation.png",
        "图3 PHEME不同观察窗口与多模态消融结果",
        width_cm=15.0,
    )
    add_body(
        doc,
        reference_element,
        "如图3所示，观察窗口从60 min延长至180 min和360 min时，MAPE由0.2388依次下降至0.1261和0.0824，表明更多早期传播轨迹能够显著减少规模预测不确定性。在180 min窗口下，删除时序模态使MAPE上升至0.1345，是各单模态消融中最明显的退化；删除文本或拓扑仅造成小幅变化，删除用户画像后MAPE略有下降。该现象说明PHEME任务的主要可预测信号来自传播增长序列，而当前用户统计量与其他模态存在一定冗余。综上所述，多模态框架能够提供稳定预测，但不同模态的边际贡献具有数据集与字段完整性依赖。",
    )
    add_body(
        doc,
        reference_element,
        "为进一步检验传播动能因子分解，本文在PHEME 180 min任务上比较V1、多变量VAE和内容—动力解耦VAE。结果如表4所示。",
    )
    add_table_title(doc, reference_element, "表4 PHEME传播动能因子分解结果")
    add_table(
        doc,
        reference_element,
        ["方法", "种子数", "测试MAPE", "MAE", "R²", "结果定位"],
        [
            ["V1 + MiniLM", "5", "0.1247±0.0019", "2.9919", "0.7392", "多模态参考模型"],
            ["V2/VAE因子", "1", "0.1231", "2.9521", "0.7436", "验证集选择的最优单种子"],
            ["V2/解耦+匹配置换", "5", "0.1243±0.0015", "3.0009", "0.7375", "解释性与稳健性分支"],
        ],
        [3.8, 1.5, 2.6, 2.1, 1.8, 4.1],
        font_size=8.3,
        left_columns={0, 5},
    )
    add_body(
        doc,
        reference_element,
        "验证集敏感性分析选择K=4和KL权重0.1。该配置在seed=42下取得0.1231的测试MAPE，四个潜变量均保持活跃，说明融合表征中存在可被低维变量概括的传播信息。需要指出的是，该结果来自单个随机种子，其1.30%的相对改进不能直接解释为稳定的总体优势。",
    )
    add_figure(
        doc,
        reference_element,
        "fig6_v2_c1_latent_factors.png",
        "图4 V2传播动能潜变量分布及其与级联增长的关系",
        width_cm=15.0,
    )
    add_body(
        doc,
        reference_element,
        "图4展示了潜变量的低维投影和增长相关性。不同因子在样本空间中形成连续而非完全离散的结构，表明传播动能更适合作为多因素共同作用的连续状态，而不是单一类别标签。为减少内容因素与动态因素之间的混杂，本文进一步将文本编码映射为内容因子，将拓扑、时序和用户表示映射为动力因子，并利用目标匹配的文本置换构造反事实压力测试。",
    )
    add_figure(
        doc,
        reference_element,
        "fig8_v2_disentangled_multiseed.png",
        "图5 解耦传播动能模型的多随机种子性能与反事实压力测试",
        width_cm=15.0,
    )
    add_body(
        doc,
        reference_element,
        "由图5可知，解耦模型在五个随机种子下的测试MAPE为0.1243±0.0015，较V1五种子均值降低0.32%，并在30%文本特征扰动下保持0.1243的MAPE。然而，目标匹配文本置换使MAPE上升至0.1263，说明内容替换仍会改变部分预测。综上所述，当前证据支持“低维传播因子具有解释与稳健性价值”，但不足以将反事实约束表述为严格的因果识别证明。",
    )

    add_heading(doc, reference_element, "5.2.2 破圈临界预警有效性分析", 4)
    add_body(
        doc,
        reference_element,
        "极化圈层中的传播规模可能在较长时间内保持平稳，而跨社区边和图频谱能量已开始积聚。为验证频谱—动态图协同预警模块，本文将完整C2模型与静态图、动态趋势、仅动态图、仅社区特征以及删除时序趋势、低频能量和跨社区拓扑的变体进行比较。主实验在五个随机种子下报告结果，完整模型的表现如表5所示。",
    )
    add_table_title(doc, reference_element, "表5 HeteroRumorDyn-C2破圈预警结果")
    add_table(
        doc,
        reference_element,
        ["数据集", "AUC", "F1", "Recall@10%", "平均提前时间/min", "最优对照及AUC"],
        [
            ["PHEME", "0.8433±0.0008", "0.7842±0.0005", "0.1793±0.0046", "185.65±8.65", "w/o low-freq: 0.8430"],
            ["Twitter15", "0.9571±0.0003", "0.9425±0.0048", "0.1266", "180.51±15.40", "w/o cross-comm: 0.9636"],
            ["Twitter16", "0.9653±0.0019", "0.9181±0.0032", "0.1491", "164.86", "Dynamic trend: 0.9666"],
        ],
        [2.0, 2.3, 2.3, 2.2, 2.6, 4.3],
        font_size=8.1,
        left_columns={5},
    )
    add_figure(
        doc,
        reference_element,
        "fig10_c2_breakout_multiseed.png",
        "图6 C2破圈预警的多随机种子结果与模块消融",
        width_cm=15.2,
    )
    add_body(
        doc,
        reference_element,
        "由表5和图6可知，完整C2模型在三个主要数据集上均获得较高且波动较小的AUC，说明窗口级动态、频谱与社区信息能够形成稳定的破圈风险排序。在PHEME上，完整模型的AUC为0.8433，略高于删除低频能量后的0.8430和删除时序趋势后的0.8426；删除跨社区拓扑后AUC下降至0.8340，平均提前时间由185.65 min下降至约160.04 min。这表明PHEME中的跨社区桥接变化对提前预警具有更直接的作用，而低频能量提供的增益较小。",
    )
    add_body(
        doc,
        reference_element,
        "在Twitter15上，删除跨社区特征后的AUC达到0.9636，高于完整模型的0.9571；在Twitter16上，动态趋势基线以0.9666略高于完整模型的0.9653，但删除跨社区特征会使AUC显著下降至0.9482。两组结果说明社区代理的可靠性具有数据集依赖性：Twitter15的分支社区可能引入冗余或噪声，而Twitter16的跨社区变化则包含重要判别信息。因此，低频能量和社区拓扑不能被表述为在所有数据集上均产生单调增益，其作用取决于传播树质量、社区构造方式和事件规模。",
    )
    add_body(
        doc,
        reference_element,
        "在按时间排序的seed=42压力测试中，C2在PHEME、Twitter15和Twitter16上的AUC分别达到0.8903、0.9711和0.9741，表明风险排序在时间分布变化下仍保持较强区分能力。综上所述，C2模块验证了动态趋势与跨社区信号对破圈预警的总体价值，同时也揭示了当前启发式社区与破圈标签需要进一步校准的边界。",
    )

    add_heading(doc, reference_element, "5.2.3 时滞闭环控制有效性分析", 4)
    add_body(
        doc,
        reference_element,
        "为验证破圈风险驱动的时滞闭环控制模块，本文使用C2输出的风险分数和观测级联快照进行控制仿真，并与固定干预、影响力阻断、随机同预算、固定同预算及ED-ID-adapted策略比较。为避免不同资源投入造成不公平比较，主要结论基于相同平均干预成本下的结果。",
    )
    add_table_title(doc, reference_element, "表6 同预算条件下闭环控制结果")
    add_table(
        doc,
        reference_element,
        ["数据集", "随机同预算抑制率", "固定同预算抑制率", "ED-ID同预算抑制率", "本文抑制率", "本文成本", "本文收益成本比"],
        [
            ["PHEME", "0.0346±0.0013", "0.0127±0.0001", "0.0360±0.0002", "0.0485±0.0004", "0.2619±0.0019", "4.9119±0.0133"],
            ["Twitter15", "0.1161±0.0035", "0.0411±0.0015", "0.1031±0.0043", "0.1455±0.0058", "0.4664±0.0171", "16.9820±0.1351"],
            ["Twitter16", "0.1110±0.0036", "0.0402±0.0001", "0.1033±0.0004", "0.1445±0.0003", "0.4482±0.0014", "16.6885±0.1082"],
        ],
        [1.8, 2.2, 2.2, 2.2, 2.2, 2.3, 2.7],
        font_size=7.6,
    )
    add_figure(
        doc,
        reference_element,
        "fig11_c3_control_multiseed.png",
        "图7 C3闭环控制的抑制率、成本与收益成本比",
        width_cm=15.2,
    )
    add_body(
        doc,
        reference_element,
        "由表6和图7可知，在相同预算下，HeteroRumorDyn-C3在PHEME、Twitter15和Twitter16上的传播抑制率分别为0.0485、0.1455和0.1445。相对于随机同预算策略，抑制率分别提高约40.2%、25.3%和30.2%；相对于ED-ID同预算策略，分别提高约34.7%、41.1%和39.9%。这说明性能提升主要来自风险感知的时机选择和自适应脉冲强度，而不是简单增加干预预算。",
    )
    add_body(
        doc,
        reference_element,
        "博弈与触发机制的消融进一步揭示了两类贡献。删除主从博弈后，三个数据集的抑制率分别降至0.0370、0.0819和0.0796，说明传播者规避响应会显著影响最优控制强度；删除事件触发后，策略在所有样本上连续干预，成本上升至1.2，而抑制率仅为0.0435、0.0793和0.0807，收益成本比明显下降。由此可见，博弈模块主要改善动作强度，事件触发机制主要减少无效干预并提高资源利用效率。",
    )
    add_body(
        doc,
        reference_element,
        "需要注意的是，ED-ID-adapted在PHEME原生预算下取得0.0591的抑制率，高于本文模型的0.0485，但其平均成本为0.4292，高于本文的0.2619；在Twitter15上二者抑制率相当，而本文成本更低。该现象表明控制策略应同时报告效果与成本，不能仅依据抑制率判断优劣。时间切分实验中，本文模型在PHEME、Twitter15和Twitter16上的抑制率分别为0.0275、0.1692和0.1965，说明闭环策略在时间分布变化下仍能发挥作用，但不同数据集的可控空间存在明显差异。",
    )

    add_heading(doc, reference_element, "5.2.4 综合性能与案例分析", 4)
    add_body(
        doc,
        reference_element,
        "前述实验分别验证了传播动力辨识、破圈预警和闭环控制。为考察三个模块在单个级联中的衔接关系，本文进一步选取预警成功、高控制收益和误报挑战三类案例，分析风险轨迹、预警时间和控制结果，如图8所示。",
    )
    add_figure(
        doc,
        reference_element,
        "fig12_c2_c3_case_studies.png",
        "图8 破圈预警与闭环控制典型案例",
        width_cm=15.2,
    )
    add_body(
        doc,
        reference_element,
        "在PHEME案例544438653410639872中，模型在真实破圈前360 min给出风险告警，风险得分为0.7686，闭环控制取得0.1701的抑制率，而同预算随机策略未产生有效抑制。这说明提前风险信号能够为后续控制保留足够时间。在Twitter15案例499679379820412928中，风险得分达到0.9818，模型提前60 min告警并实现0.4348的抑制率，同预算随机策略同样未取得明显效果，进一步表明高风险窗口的定向干预具有实际价值。",
    )
    add_body(
        doc,
        reference_element,
        "Twitter16案例638047610973089793属于未破圈样本，但模型给出0.9232的高风险得分并触发干预，本文策略与同预算随机策略的抑制率分别为0.2275和0.2283。该误报案例表明，当传播结构表现出强增长与跨社区特征但最终未超过启发式破圈阈值时，模型可能实施不必要的干预。其原因既可能是风险模型的校准不足，也可能是二值破圈标签无法完整表达“高风险但未越界”的连续状态。",
    )
    add_body(
        doc,
        reference_element,
        "综上所述，实验形成了“早期传播表征—破圈风险排序—风险驱动控制”的完整证据链：C1在跨平台规模预测上具有较强优势并支持低维因子分解；C2能够稳定识别高风险级联，但频谱和社区特征的贡献受数据质量影响；C3在同预算条件下显著提高控制效率。与此同时，反事实因果解释、启发式破圈标签和仿真控制仍限定了当前结论的外推范围。",
    )

    add_heading(doc, reference_element, "6 总结", 2)
    add_body(
        doc,
        reference_element,
        "本文针对异构社交生态中传播动能难辨识、极化圈层破圈信号难提前感知以及博弈时滞条件下干预策略易失准的问题，提出HeteroRumorDyn谣言传播动力与控制模型。该模型通过多模态异构图编码与变分因子分解刻画传播动能，利用时序趋势、动态图信息、低频能量和跨社区拓扑构建破圈风险，并将风险反馈引入Stackelberg主从博弈与事件触发脉冲控制。实验结果表明，本文模型在Twitter15和Twitter16的180 min传播规模预测中显著降低MAPE，在三个主要数据集上获得稳定的破圈风险排序，并在同预算条件下较随机、固定和ED-ID改编策略取得更高传播抑制率。多随机种子、时间切分和案例分析进一步验证了方法的稳定性，同时揭示了不同数据集上社区与频谱信息贡献的差异。",
    )
    add_body(
        doc,
        reference_element,
        "当前研究仍存在若干边界。首先，破圈事件与社区结构主要由传播分支和跨社区比例启发式构造，C2当前采用窗口级特征分类器验证机制变量，尚未完成真实社交关系上的端到端TGN训练；其次，Weibo数据缺少真实时间戳、源文本和用户画像，其结果只能作为代理证据；再次，C3基于风险分数和级联快照进行控制仿真，尚未获得真实平台干预日志；最后，内容置换压力测试支持表征稳健性，但尚不足以证明严格的因果可识别性。未来将引入具有真实时间、用户关系和干预记录的数据，联合学习社区结构与破圈标签，完善端到端动态图预警模型，并结合结构因果模型、分布外验证和在线策略评估，构建从传播动能辨识、临界预警到可验证干预效果的闭环谣言治理体系。",
    )

    append_reference(
        doc,
        "Song, X., Hu, M., Zhou, F. & Dai, L. CHGNN: A Semi-Supervised Contrastive Hypergraph Learning Network. arXiv preprint arXiv:2312.04452 (2023).",
    )
    append_reference(
        doc,
        "Liu, Y., Yu, S., Wen, W. & Han, X. MCDAN: A Multi-scale Context-enhanced Dynamic Attention Network for Diffusion Prediction. arXiv preprint arXiv:2303.06213 (2023).",
    )
    append_reference(
        doc,
        "Jin, F., Dougherty, E., Saraf, P., Cao, Y. & Ramakrishnan, N. Epidemiological Modeling of News and Rumors on Twitter. Proceedings of the 7th Workshop on Social Network Mining and Analysis, 1-9 (2013). https://doi.org/10.1145/2501025.2501027.",
    )
    append_reference(
        doc,
        "Jiang, Y., Chen, X., Hua, Q.-S., Gong, M., Wu, W., Zhang, J. & Wu, J. Evolutionary Dynamics of Information Diffusion over Social Networks. arXiv preprint arXiv:2407.04861 (2024).",
    )

    new_body_children = list(doc._element.body)
    new_chapter5_index = new_body_children.index(chapter5._p)
    after_digest = element_digest(new_body_children[: new_chapter5_index + 1])
    if before_digest != after_digest:
        raise RuntimeError("第1—4章及第5章标题的OOXML发生变化，已中止保存。")

    doc.save(OUTPUT)
    print(f"saved={OUTPUT}")
    print(f"unchanged_prefix_sha256={before_digest}")


if __name__ == "__main__":
    build()
