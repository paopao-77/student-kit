from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


BASE = Path(__file__).resolve().parents[1]
SOURCE = BASE / "论文写作" / "V6.1_raw_weibo_updated_baseline_fixed.docx"
OUTPUT = BASE / "论文写作" / "V6.1_final_strategy_applied.docx"


def set_text(paragraph, text):
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def set_cell(cell, text):
    set_text(cell.paragraphs[0], text)
    for p in cell.paragraphs[1:]:
        set_text(p, "")


def set_row(row, values):
    for cell, value in zip(row.cells, values):
        set_cell(cell, value)


def insert_paragraph_after(paragraph, text="", style=None):
    new_p = paragraph.insert_paragraph_before(text)
    paragraph._p.addnext(new_p._p)
    if style:
        new_p.style = style
    return new_p


def insert_table_after(doc, paragraph, rows):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    if doc.tables:
        table.style = doc.tables[1].style
    for r_idx, row_vals in enumerate(rows):
        set_row(table.rows[r_idx], row_vals)
        for cell in table.rows[r_idx].cells:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if r_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
    paragraph._p.addnext(table._tbl)
    return table


def main():
    doc = Document(SOURCE)

    replacements = {
        2: "社交媒体实时化与圈层化使谣言扩散呈现欺骗性耦合、隐性积聚与对抗反馈等复杂动态。核心挑战在于：如何在多源异构信号不完备、极化拓扑可能突然破圈、平台干预存在时滞的条件下，学习能够支撑风险预警的传播动能表征。针对该问题，本文提出面向异构社交生态的HeteroRumorDyn框架。该框架首先融合文本语义、用户属性、异构拓扑与历史传播序列，并通过变分因子分解和反事实一致性约束学习稳健、可解释的传播动能因子；随后，将传播动能因子与时序趋势、动态图记忆、低频传播能量和跨社区桥接变化联合建模，用于识别由局部扩散转向跨圈级联的破圈风险；最后，将风险信号接入时滞传播与主从博弈框架，构建事件触发的闭环干预仿真机制，以评估风险感知策略在有限干预成本下的潜在控制效果。实验结果表明，HeteroRumorDyn在传播规模预测、破圈风险排序和同预算控制仿真中整体优于可比基线；同时，结果也显示潜变量分解、频谱特征和社区桥接信号的增益具有数据集依赖性。本文因此将贡献界定为传播动能表征与破圈风险预警框架，并将控制结果作为下游仿真验证，而非真实平台干预效果证明。",
        5: "Real-time and clustered social media ecosystems make rumor diffusion exhibit deceptive coupling, latent accumulation and adversarial feedback. The central challenge is to learn propagation-dynamics representations that support circle-breaking risk warning under incomplete heterogeneous signals, polarized topology and delayed platform intervention. We propose HeteroRumorDyn, a framework that first learns robust and interpretable propagation-dynamics factors by fusing textual, user, topological and temporal signals with variational factorization and counterfactual consistency constraints. It then combines these factors with temporal trends, dynamic graph memory, low-frequency propagation energy and cross-community bridging signals to estimate the risk that local diffusion turns into cross-circle cascades. Finally, the risk signal is connected to a delayed Stackelberg intervention simulator with event-triggered impulses to evaluate downstream control value under finite budgets. Experiments show that HeteroRumorDyn improves cascade-size prediction, breakout risk ranking and same-budget control simulation over comparable baselines, while also revealing dataset-dependent gains of latent factors, spectral features and community proxies.",
        17: "针对上述问题，本文提出面向异构社交生态的谣言传播动力与控制模型（HeteroRumorDyn），以传播动能表征支撑破圈风险预警，并进一步通过控制仿真检验风险信号的下游干预价值。本文的主要贡献如下：",
        18: "1、提出一种面向异构传播生态的传播动能表征方法。该方法融合文本语义、用户行为、异构拓扑与历史传播信息，并通过变分因子分解与反事实一致性约束学习低维传播动能因子，从而提升多源传播表征的稳健性与可解释性。",
        19: "2、提出一种传播动能驱动的破圈风险预警方法。该方法将传播动能因子与时序趋势、动态图记忆、低频传播能量和跨社区拓扑变化联合建模，用于识别由局部扩散转向跨圈级联的临界信号，并输出下一时刻破圈风险与传播状态。",
        20: "3、构建一种风险驱动的时滞闭环控制仿真框架。该框架将破圈风险作为反馈信号，引入主从博弈和事件触发脉冲机制，在同预算条件下评估动态干预策略的抑制效果与成本效率。",
        21: "实验结果表明，该模型能够形成有效的传播动能表征，提升破圈风险排序与提前预警能力，并在控制仿真中体现风险信号的下游应用价值。与此同时，本文不将反事实一致性等同于严格因果识别，也不将仿真结果表述为真实平台控制效果。",
        63: "HeteroRumorDyn的核心目标是学习可解释的传播动能表示，并利用该表示提升破圈风险预警能力；在此基础上，本文进一步构建风险驱动的闭环控制仿真模块，用于验证预警信号在下游动态干预中的应用价值。如图2所示，该模型由三个顺序衔接的部分构成。首先，传播动能表征模块联合编码消息语义、用户属性、异构拓扑与历史传播序列，并通过变分因子分解和反事实一致性约束削弱欺骗性耦合噪声。其次，频谱—动态图协同预警模块利用时序趋势、动态图记忆、低频能量和跨社区桥接变化识别破圈前兆。最后，时滞闭环控制仿真模块将平台干预与传播者规避建模为Stackelberg主从博弈，并依据破圈风险实施事件触发脉冲控制，以比较不同策略在同预算条件下的相对抑制能力。三个部分依次形成“动能表征—风险预警—控制仿真验证”的建模流程，而不是三个彼此独立的平权理论任务。在数据接入层面，模型使用模态掩码记录文本、拓扑、时序和用户画像的可用性；当数据仅提供事件顺序而缺少可比真实时间戳时，以事件序窗口替代分钟窗口，并在实验解释中单独限定其外推范围。",
        66: "4.1基于反事实一致性的传播动能表征模块",
        67: "异构社交网络中的消息语义、用户行为与传播拓扑往往受到伪信誉、强关联和高煽动性内容的共同影响，直接拼接多源特征容易把混淆因素误判为真实传播动力。为此，本文将异构图编码、变分因子分解与反事实一致性约束统一到传播动能表征过程中，使模型学习对非关键扰动更稳定、对真实扩散状态更敏感的低维动能表征，如图3所示。",
        81: "其中，传播状态预测损失约束规模与用户参与状态预测，重构损失保持多源信息可重构性，KL项使潜变量分布平滑，反事实一致性项抑制非关键扰动造成的预测漂移。通过上述联合优化，该模块不再只是多源特征的简单压缩，而是面向传播动力分析的稳健动能表征。需要强调的是，本文的反事实扰动用于提高表征稳健性与解释性，并不等同于对真实因果结构的严格可识别性证明。",
        83: "极化圈层内的谣言可能长期保持较小的显性规模，但其传播信号会在低频成分和跨社区拓扑中持续积聚。仅依赖静态图或单一规模阈值难以区分一般波动与破圈前兆。受多视图时空动态图学习方法[16]启发，本文联合建模历史传播序列、动态图记忆、图频谱能量和跨社区桥接变化，以识别由局部扩散转向跨圈级联的临界状态。与仅依赖传播规模变化的预警方法不同，本文将低频能量用于刻画圈层内隐性积聚，将跨社区桥接强度用于刻画破圈通道激活，并将传播动能因子作为上游统一表征输入风险预测器，从而实现从“动能积累”到“结构突破”的联合预警。",
        97: "4.3破圈风险驱动的时滞闭环控制仿真模块",
        98: "由于真实平台干预日志通常难以获得，本文将C3设计为风险驱动的闭环控制仿真模块，用于评估C2输出的风险信号在动态干预场景中的应用价值。该模块不直接声称真实平台控制效果，而是在统一预算约束下比较不同干预策略的相对抑制能力与成本效率。平台干预通常需要经历识别、审核、策略下发和生效过程，而传播者可能通过内容改写或路径迁移进行规避，因此干预动作与传播状态之间存在显著时滞。本文结合实时影响阻断估计思想[17]，以破圈风险为反馈信号，将时滞传播方程、Stackelberg主从博弈与事件触发脉冲控制统一到仿真框架中。",
        115: "HeteroRumorDyn的输入包括异构传播图、历史传播序列、平台反馈、社区桥接状态和模态覆盖标记，输出包括传播动能因子、下一时刻破圈风险、传播状态以及仿真环境下的动态干预策略。模型按照时间窗口滚动执行：先完成多源编码与传播动能表征，再更新动态图记忆和频谱风险信号，最后求解平台—传播者主从博弈并执行事件触发控制仿真。具体过程如算法所示。",
        119: "本章围绕第4章提出的传播动能表征、破圈临界预警和时滞闭环控制仿真三个部分开展实验验证。实验首先说明数据来源、对比方法、评价指标、参数设置和可复现性协议，随后按照“主张—证据”对应的方式分析传播规模预测、传播动能因子、破圈预警与控制仿真的有效性，并通过多随机种子、时间切分、外部留出、消融和典型案例检验结果的稳定性与适用边界。",
        137: "本节按照第4章的证据链组织实验。首先检验多模态传播规模预测与传播动能因子分解，其次分析低频能量和跨社区拓扑对破圈预警的贡献，最后比较风险驱动闭环控制仿真与固定、随机及演化博弈策略的抑制效果和成本。对于增益不稳定或特征贡献非单调的结果，本文在相应小节中正面解释其数据依赖性与适用边界。",
        138: "5.2.1 传播动能表征有效性分析",
        139: "传播动能表征模块需要从有限观察窗口内的文本、拓扑、时序和用户信息中预测最终级联规模，并将多源表征压缩为可解释的低维因子。为验证该模块的有效性，本文首先在统一的180 min观察窗口和相同测试样本上，将HeteroRumorDyn-V1与MIDPMS、DSHCL、Inf-VAE和CD-SEIZ进行比较，结果如表4所示。raw-Weibo使用原始直连接入口径，本文将其作为模型内补充验证在表5报告，不与表4中PHEME、Twitter15和Twitter16的同任务基线混列。",
        140: "表4 180 min观察窗口下传播规模预测结果",
        142: "由表4可知，HeteroRumorDyn-V1在Twitter15和Twitter16上取得最低MAPE，分别为0.1785和0.2010；相对于对应数据集表现最好的同领域基线，误差分别降低约28.0%和23.7%。新增Inf-VAE-adapted在PHEME、Twitter15和Twitter16上的MAPE分别为0.2084±0.0115、0.2604±0.0087和0.2750±0.0265，整体优于或接近MIDPMS/DSHCL，但仍弱于本文模型，说明通用变分潜变量扩散表示难以完全替代面向谣言传播场景设计的多模态异构动能建模。在PHEME上，CD-SEIZ的MAPE为0.1234，略低于本文模型的0.1247，但二者差距较小，配对自助法置信区间覆盖零。这说明PHEME中较短、较规则的会话级联能够被经典动力学较好拟合，而在Twitter15/16这类结构差异更明显的传播过程中，多模态表示具有更稳定的应用价值。",
        147: "为进一步检验传播动能因子分解，本文在PHEME 180 min任务和raw-Weibo 180事件任务上比较V1、多变量VAE和内容—动力解耦VAE。结果如表5所示。",
        148: "表5 PHEME与raw-Weibo传播动能因子分解结果",
        157: "极化圈层中的传播规模可能在较长时间内保持平稳，而跨社区边和图频谱能量已开始积聚。为验证频谱—动态图协同预警模块，本文将完整C2模型与静态图、动态趋势、仅动态图、仅社区特征以及删除时序趋势、低频能量和跨社区拓扑的变体进行比较。主实验在五个随机种子下报告结果，完整模型的表现如表6所示；Weibo采用事件序窗口，主报告口径为order_window_size=50。",
        158: "表6 HeteroRumorDyn-C2破圈预警结果",
        161: "由表6和图6可知，完整C2模型在PHEME、Twitter15、Twitter16和raw-Weibo上均获得较高且波动较小的AUC，说明窗口级动态、频谱与社区信息能够形成稳定的破圈风险排序。在PHEME上，完整模型的AUC为0.8433，略高于删除低频能量后的0.8430和删除时序趋势后的0.8426；删除跨社区拓扑后AUC下降至0.8340，平均提前时间由185.65 min下降至约160.04 min。这表明PHEME中的跨社区桥接变化对提前预警具有更直接的作用，而低频能量提供的增益较小。",
        165: "为验证破圈风险信号的下游干预价值，本文使用C2输出的风险分数和观测级联快照进行控制仿真，并与固定干预、影响力阻断、随机同预算、固定同预算及ED-ID-adapted策略比较。为避免不同资源投入造成不公平比较，主要结论基于相同平均干预成本下的结果；原生预算结果仅用于分析效果—成本权衡。",
        166: "表7 同预算条件下闭环控制结果",
        168: "图7 C3闭环控制仿真的抑制率、成本与收益成本比",
        169: "由表7和图7可知，在相同预算下，HeteroRumorDyn-C3在PHEME、Twitter15、Twitter16和raw-Weibo上的传播抑制率分别为0.0485、0.1455、0.1445和0.2304。相对于随机同预算策略，前三个数据集的抑制率分别提高约40.2%、25.3%和30.2%；相对于ED-ID同预算策略，分别提高约34.7%、41.1%和39.9%。在raw-Weibo上，抑制率较随机同预算、固定同预算和ED-ID同预算分别提高约31.8%、266.2%和59.6%。这说明性能提升主要来自风险感知的时机选择和自适应脉冲强度，而不是简单增加干预预算。",
        170: "博弈与触发机制的消融进一步揭示了两类贡献。删除主从博弈后，PHEME、Twitter15和Twitter16的抑制率分别降至0.0370、0.0819和0.0796，说明传播者规避响应会显著影响控制强度；删除事件触发后，策略在所有样本上连续干预，成本上升至1.2，而抑制率仅为0.0435、0.0793和0.0807，收益成本比明显下降。在raw-Weibo上，删除主从博弈和删除事件触发后的抑制率分别降至0.1255和0.1294，同样低于完整C3。由此可见，博弈模块主要改善动作强度，事件触发机制主要减少无效干预并提高资源利用效率。",
        173: "前述实验分别验证了传播动能表征、破圈预警和闭环控制仿真。为考察三个模块在单个级联中的衔接关系，本文进一步选取预警成功、高控制收益和误报挑战三类案例，分析风险轨迹、预警时间和控制结果，如图8所示。",
        178: "综上所述，实验形成了“早期传播表征—破圈风险排序—风险驱动控制仿真”的证据链：C1在跨平台规模预测上具有较强优势，并在raw-Weibo上确认文本和用户画像接入后的中文级联预测可行性；V2/C1支持低维因子分解，但Weibo上的增益不稳定；C2能够稳定识别高风险级联，但频谱和社区特征的贡献受数据质量影响；C3在同预算仿真条件下提高控制效率，说明风险信号具有下游干预应用潜力。与此同时，反事实因果解释、启发式破圈标签、事件序时间代理和无真实干预日志仍限定了当前结论的外推范围。",
        180: "本文针对异构社交生态中传播动能难表征、极化圈层破圈信号难提前感知以及平台干预存在时滞的问题，提出HeteroRumorDyn谣言传播动力与控制仿真框架。该框架通过多模态异构图编码与变分因子分解刻画传播动能，利用时序趋势、动态图信息、低频能量和跨社区拓扑构建破圈风险，并将风险反馈接入Stackelberg主从博弈与事件触发脉冲控制仿真。实验结果表明，本文模型在Twitter15和Twitter16的180 min传播规模预测中显著降低MAPE，在raw-Weibo原始数据直连接口下取得0.1477的V1 MAPE、0.9706的C2 AUC和0.2304的C3同预算抑制率，在主要数据集上获得稳定的破圈风险排序，并在同预算仿真条件下较随机、固定和ED-ID改编策略取得更高传播抑制率。多随机种子、时间切分、外部留出和案例分析进一步验证了方法的稳定性，同时揭示了不同数据集上社区与频谱信息贡献的差异。",
        181: "当前研究仍存在若干边界。首先，破圈事件与社区结构主要由传播分支和跨社区比例启发式构造，C2当前采用窗口级特征分类器验证机制变量，尚未完成真实社交关系上的端到端TGN训练；其次，raw-Weibo虽然已经接入源文本和用户画像，但仍缺少可与PHEME直接对齐的真实时间戳和真实社交关系，因而其提前量和社区解释需要按事件序与分支代理口径理解；再次，C3基于风险分数和级联快照进行控制仿真，尚未获得真实平台干预日志，不能被解释为真实平台干预效果证明；最后，内容置换压力测试支持表征稳健性，但尚不足以证明严格的因果可识别性。未来将引入具有真实时间、用户关系和干预记录的数据，联合学习社区结构与破圈标签，完善端到端动态图预警模型，并结合结构因果模型、分布外验证和在线策略评估，构建从传播动能表征、临界预警到可验证干预效果的闭环谣言治理体系。",
    }

    for idx, text in replacements.items():
        set_text(doc.paragraphs[idx], text)

    # Insert reproducibility protocol and evidence table after the parameter note.
    anchor = doc.paragraphs[135]
    p = insert_paragraph_after(anchor, "5.1.5 可复现性与报告协议", style="Heading 3")
    p2 = insert_paragraph_after(
        p,
        "为避免不同数据接入口径造成结果不可比，本文将raw-Weibo原始直连接口作为主报告口径，历史接入口径仅作为实现参考，不参与主表对比。所有主实验均采用固定数据划分、五个随机种子和统一窗口参数，并在实验索引中记录预处理脚本、配置文件、运行命令和结果入口。Weibo实验采用180个事件观察窗口，C2/C3主报告口径采用order_window_size=50；社区划分与破圈标签构造规则在结果索引和可复现性清单中保持一致。",
        style="Normal",
    )
    p3 = insert_paragraph_after(p2, "表3 各模块实验验证协议与可信度说明", style="Normal")
    evidence_table = [
        ["模块", "验证目标", "数据集", "随机种子", "主指标", "稳定性检验", "当前边界"],
        ["C1", "动能表征与规模预测", "四个数据集", "5", "MAPE/MAE/R²", "时间切分、外部留出、CI", "反事实仅支持稳健性"],
        ["C2", "破圈风险预警", "四个数据集", "5", "AUC/F1/提前量", "消融、时间切分、外部留出", "社区与标签含代理"],
        ["C3", "闭环控制仿真", "四个数据集", "5", "抑制率/成本/收益成本比", "同预算对比、消融", "无真实干预日志"],
    ]
    insert_table_after(doc, p3, evidence_table)

    # Insert non-monotonic results section before the conclusion.
    anchor = next(p for p in doc.paragraphs if p.text.startswith("综上所述，实验形成了"))
    h = insert_paragraph_after(anchor, "5.2.5 非单调结果与适用边界分析", style="Heading 3")
    p1 = insert_paragraph_after(
        h,
        "首先，V2/C1在raw-Weibo上相对V1的MAPE改进极小，说明潜变量因子分解的主要价值在于解释性、稳定性和低维传播状态刻画，而不是在所有数据集上稳定提升预测精度。本文因此将V2/C1定位为传播动能解释分支，而不将其表述为普遍优于V1的预测模型。",
        style="Normal",
    )
    p2 = insert_paragraph_after(
        p1,
        "其次，C2中低频能量和跨社区特征的贡献依赖于社区划分质量、传播树结构和破圈标签定义。在Twitter15、Twitter16和raw-Weibo上，个别删除项或动态趋势基线的AUC可能接近甚至略高于完整模型。该现象说明，低频能量与社区桥接信号不是在所有数据集上均产生单调增益的黑箱组件，而是用于融合多类破圈前兆信号的风险排序框架。",
        style="Normal",
    )
    insert_paragraph_after(
        p2,
        "最后，C3的主结论基于同预算仿真比较。在原生预算下，ED-ID在个别数据集上可能达到接近甚至更高的抑制率，但其平均成本也更高。因此，本文同时报告抑制率、干预成本和收益成本比，并将同预算设置作为主要比较口径，以避免将高资源投入误解释为策略本身更优。",
        style="Normal",
    )

    doc.save(OUTPUT)
    print(f"source={SOURCE}")
    print(f"output={OUTPUT}")


if __name__ == "__main__":
    main()
