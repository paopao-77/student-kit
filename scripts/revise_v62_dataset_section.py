from pathlib import Path

from docx import Document
from docx.shared import Pt


ROOT = Path(r"F:/student-kit")
DOC_DIR = ROOT / "\u8bba\u6587\u5199\u4f5c"
INPUT_DOCX = DOC_DIR / "V6.2.docx"
OUTPUT_DOCX = DOC_DIR / "V6.2_dataset_revised.docx"

DATASET_TEXT = (
    "本文使用四份公开谣言传播数据集进行实验，包括PHEME[23]、新浪微博数据集[24]以及"
    "Twitter15和Twitter16[25,26]。数据集包含源帖或推文文本、回复或转发结构、用户公开属性"
    "与时间信息等多维属性，能够覆盖真实时间演化、跨平台传播结构和中文大规模级联场景。"
    "为确保数据的多样性与代表性，本文分别在英文事件会话、Twitter级联和新浪微博原始级联上"
    "验证模型的传播规模预测、破圈预警与控制仿真能力。其中，Weibo基于原始文件重新接入，"
    "在180个事件观察窗口下文本覆盖率为98.14%，用户画像覆盖率为73.16%；其提前量按事件序单位解释。"
    "各数据集统计如表1所示。"
)

DATASET_REFS = [
    (
        "Zubiaga, A., Liakata, M., Procter, R., Hoi, G. W. S. & Tolmie, P. "
        "Analysing how people orient to and spread rumours in social media by looking at conversational threads. "
        "PLoS ONE 11(3), e0150989 (2016). https://doi.org/10.1371/journal.pone.0150989."
    ),
    (
        "Ma, J., Gao, W., Wei, Z., Lu, Y. & Wong, K.-F. "
        "Detect Rumors Using Time Series of Social Context Information on Microblogging Websites. "
        "Proceedings of the 24th ACM International on Conference on Information and Knowledge Management, "
        "1751-1754 (2015). https://doi.org/10.1145/2806416.2806607."
    ),
    (
        "Ma, J., Gao, W., Mitra, P., Kwon, S., Jansen, B. J., Wong, K.-F. & Cha, M. "
        "Detecting Rumors from Microblogs with Recurrent Neural Networks. "
        "Proceedings of the Twenty-Fifth International Joint Conference on Artificial Intelligence, "
        "3818-3824 (2016)."
    ),
    (
        "Ma, J., Gao, W. & Wong, K.-F. "
        "Detect Rumors in Microblog Posts Using Propagation Structure via Kernel Learning. "
        "Proceedings of the 55th Annual Meeting of the Association for Computational Linguistics, "
        "708-717 (2017). https://doi.org/10.18653/v1/P17-1066."
    ),
]


def set_paragraph_text(paragraph, text: str) -> None:
    paragraph.text = text
    for run in paragraph.runs:
        run.font.name = "SimSun"
        run.font.size = Pt(10.5)


def main() -> None:
    if not INPUT_DOCX.exists():
        raise FileNotFoundError(INPUT_DOCX)

    doc = Document(INPUT_DOCX)

    target = None
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if text.startswith("为确保实验同时覆盖真实时间演化"):
            target = paragraph
            break
    if target is None:
        raise RuntimeError("Could not locate the original dataset paragraph.")
    set_paragraph_text(target, DATASET_TEXT)

    full_text = "\n".join(p.text for p in doc.paragraphs)
    if "Analysing how people orient to and spread rumours" not in full_text:
        ref_style = doc.paragraphs[-1].style
        for ref in DATASET_REFS:
            p = doc.add_paragraph(style=ref_style)
            set_paragraph_text(p, ref)

    doc.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    main()
