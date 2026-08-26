from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph


ROOT = Path(r"F:/student-kit")
DOC_DIR = ROOT / "\u8bba\u6587\u5199\u4f5c"
INPUT_DOCX = DOC_DIR / "V6.1_final_strategy_applied.docx"
OUTPUT_DOCX = DOC_DIR / "V6.1_final_strategy_figures_added.docx"

FIG_DIR = ROOT / "results" / "figures"
FIG_E9 = FIG_DIR / "fig_weibo_raw_e9_diagnostics.png"
FIG_E12 = FIG_DIR / "fig_weibo_raw_e12_early_warning.png"


def paragraph_after(paragraph: Paragraph, text: str = "", style=None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style is not None:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def copy_body_format(source: Paragraph, target: Paragraph) -> None:
    target.style = source.style
    target.alignment = source.alignment
    target.paragraph_format.first_line_indent = source.paragraph_format.first_line_indent
    target.paragraph_format.left_indent = source.paragraph_format.left_indent
    target.paragraph_format.right_indent = source.paragraph_format.right_indent
    target.paragraph_format.space_before = source.paragraph_format.space_before
    target.paragraph_format.space_after = source.paragraph_format.space_after
    target.paragraph_format.line_spacing = source.paragraph_format.line_spacing
    target.paragraph_format.line_spacing_rule = source.paragraph_format.line_spacing_rule
    for run in target.runs:
        run.font.name = "SimSun"
        run.font.size = Pt(10.5)


def add_centered_picture_after(anchor: Paragraph, image_path: Path, caption: str) -> Paragraph:
    pic_para = paragraph_after(anchor)
    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_para.paragraph_format.space_before = Pt(6)
    pic_para.paragraph_format.space_after = Pt(3)
    pic_para.add_run().add_picture(str(image_path), width=Inches(5.9))

    cap_para = paragraph_after(pic_para, caption)
    cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_para.style = anchor.style
    cap_para.paragraph_format.space_before = Pt(0)
    cap_para.paragraph_format.space_after = Pt(6)
    for run in cap_para.runs:
        run.font.name = "SimSun"
        run.font.size = Pt(10.5)
    return cap_para


def main() -> None:
    if not INPUT_DOCX.exists():
        raise FileNotFoundError(INPUT_DOCX)
    for fig in [FIG_E9, FIG_E12]:
        if not fig.exists():
            raise FileNotFoundError(fig)

    doc = Document(INPUT_DOCX)
    if any("raw-Weibo补充验证与诊断" in p.text for p in doc.paragraphs):
        raise RuntimeError("Target subsection already exists; refusing to duplicate figures.")

    anchor = None
    for para in doc.paragraphs:
        if para.text.strip().startswith("综上所述，实验形成了"):
            anchor = para
            break
    if anchor is None:
        raise RuntimeError("Could not find experiment-summary insertion anchor.")

    heading = paragraph_after(anchor, "5.2.5 raw-Weibo补充验证与诊断")
    heading.style = anchor.style
    heading.paragraph_format.space_before = Pt(6)
    heading.paragraph_format.space_after = Pt(6)
    for run in heading.runs:
        run.bold = True
        run.font.name = "SimSun"
        run.font.size = Pt(10.5)

    body1 = paragraph_after(
        heading,
        "为进一步说明raw-Weibo主报告口径的稳定性，本文补充外部留出、事件序窗口、控制策略和运行效率诊断，结果如图9所示。可以看到，主报告口径下的V1、V2/C1、C2和C3结果在非分层外部留出中保持可解释变化；order_window_size=50在C2风险排序和C3抑制收益之间取得较好折中；同预算条件下，风险驱动控制相较随机、固定和ED-ID代理策略仍保持更高抑制收益。同时，运行时间诊断表明新增模块未引入不可接受的计算开销。",
    )
    copy_body_format(anchor, body1)

    cap9 = add_centered_picture_after(
        body1,
        FIG_E9,
        "图9 raw-Weibo主报告口径的外部留出、窗口敏感性、控制策略与效率诊断",
    )

    body2 = paragraph_after(
        cap9,
        "针对C2早预警任务，本文进一步统计raw-Weibo事件序条件下的正类预警召回率、提前量和误报率，如图10所示。五个随机种子下，正类预警召回率为0.9620±0.0008，平均提前量为1190.67±4.16个事件序单位，误报率为0.1518±0.0042。由于raw-Weibo缺少可靠真实时间戳，该提前量不解释为分钟或小时，而应理解为事件到达序列中的相对预警空间。",
    )
    copy_body_format(anchor, body2)

    add_centered_picture_after(
        body2,
        FIG_E12,
        "图10 raw-Weibo事件序条件下的早预警召回与提前量分析",
    )

    doc.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    main()
