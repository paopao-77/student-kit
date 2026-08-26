import csv
import json
from copy import deepcopy
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


INPUT_DOCX = Path("working_thesis_input.docx")
OUTPUT_DOCX = Path("前三章V5.0_第5章实验部分已补.docx")
FIGURE_PATH = Path("results/figures/fig9_rumdetect2017_text_fairness.png")

SUMMARY_DIR = Path("results/summary")
STATS_DIR = Path("data/processed")


def read_csv(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def read_json(path: Path) -> dict[str, Any]:
    with path.open("r", encoding="utf-8") as f:
        return json.load(f)


def fmt(value: Any, digits: int = 4) -> str:
    if value in (None, ""):
        return ""
    try:
        return f"{float(value):.{digits}f}"
    except (TypeError, ValueError):
        return str(value)


def set_run_font(run, size: float | None = None, bold: bool | None = None) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def format_paragraph(paragraph, size: float = 10.5, first_line: bool = True) -> None:
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.space_after = Pt(6)
    if first_line:
        paragraph.paragraph_format.first_line_indent = Pt(21)
    for run in paragraph.runs:
        set_run_font(run, size=size)


def add_paragraph(doc: Document, text: str, style: str = "Normal", *, first_line: bool = True):
    paragraph = doc.add_paragraph(style=style)
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5)
    format_paragraph(paragraph, first_line=first_line)
    return paragraph


def add_heading(doc: Document, text: str, level_style: str):
    paragraph = doc.add_paragraph(style=level_style)
    run = paragraph.add_run(text)
    set_run_font(run, size=12 if level_style == "Heading 3" else 11, bold=True)
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(6)
    return paragraph


def add_caption(doc: Document, text: str):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    set_run_font(run, size=10, bold=False)
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(6)
    return paragraph


def set_cell_margins(cell, top=90, start=90, bottom=90, end=90):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def shade_cell(cell, fill: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_widths(table, widths: list[float]) -> None:
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "999999")


def set_cell_text(cell, text: str, *, bold: bool = False, align_center: bool = False, size: float = 9.0):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if align_center else WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(str(text))
    set_run_font(run, size=size, bold=bold)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell)


def apply_table_style(table) -> None:
    for style_name in ("Table Grid", "Normal Table"):
        try:
            table.style = style_name
            return
        except KeyError:
            continue


def add_table(doc: Document, caption: str, headers: list[str], rows: list[list[str]], widths: list[float]):
    add_caption(doc, caption)
    table = doc.add_table(rows=1, cols=len(headers))
    apply_table_style(table)
    table.autofit = False
    set_table_borders(table)
    set_table_widths(table, widths)
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, align_center=True, size=8.8)
        shade_cell(table.rows[0].cells[idx], "D9EAF7")
    for data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(data):
            set_cell_text(cells[idx], value, align_center=idx != 0, size=8.6)
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    return table


def delete_paragraph(paragraph) -> None:
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None


def move_before_reference(doc: Document, new_elements: list[Any]) -> None:
    ref_para = None
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == "参考文献":
            ref_para = paragraph
            break
    if ref_para is None:
        raise ValueError("Cannot find 参考文献 heading")
    ref_node = ref_para._p
    for element in new_elements:
        ref_node.addprevious(element)


def body_insert_index(doc: Document) -> int:
    body = doc._body._element
    elements = list(body)
    sect_pr = body.sectPr
    if sect_pr is None:
        return len(elements)
    return elements.index(sect_pr)


def collect_new_elements(doc: Document, start_index: int) -> list[Any]:
    body = doc._body._element
    elements = list(body)
    sect_pr = body.sectPr
    end_index = elements.index(sect_pr) if sect_pr is not None else len(elements)
    return elements[start_index:end_index]


def remove_old_empty_experiment_body(doc: Document) -> None:
    for idx, paragraph in enumerate(list(doc.paragraphs)):
        if paragraph.text.strip() == "5 实验":
            # Remove blank paragraphs between the empty chapter heading and references.
            for nxt in list(doc.paragraphs)[idx + 1 :]:
                text = nxt.text.strip()
                if text == "参考文献":
                    return
                if text == "":
                    delete_paragraph(nxt)
            return


def build_dataset_rows() -> list[list[str]]:
    stats = {
        "PHEME": read_json(STATS_DIR / "pheme" / "stats.json"),
        "Twitter15-RD17": read_json(STATS_DIR / "twitter15_rumdetect2017" / "stats.json"),
        "Twitter16-RD17": read_json(STATS_DIR / "twitter16_rumdetect2017" / "stats.json"),
    }
    return [
        [
            name,
            str(item["num_samples"]),
            str(item.get("num_edges", "")),
            str(item.get("samples_with_source_text", "")),
            fmt(item.get("avg_cascade_nodes", ""), 2),
            "源文本、传播边、相对延迟" if name != "PHEME" else "源文本、事件文本、传播边、相对延迟、用户画像",
        ]
        for name, item in stats.items()
    ]


def build_pheme_rows() -> list[list[str]]:
    rows = read_csv(SUMMARY_DIR / "paper_v2_c1_main_table.csv")
    return [
        [
            row["method"],
            row["n_seeds"],
            row["mape_report"],
            row["mae"],
            row["rmse"],
            row["r2"],
            row["relative_mape_gain_vs_v1"],
        ]
        for row in rows
    ]


def build_fairness_rows() -> list[list[str]]:
    rows = read_csv(SUMMARY_DIR / "v1_rumdetect2017_text_fairness.csv")
    output = []
    specs = [
        ("Hash文本", "hash"),
        ("MiniLM文本", "minilm"),
        ("去除文本", "no_text"),
    ]
    for row in rows:
        dataset_label = row["dataset"].replace("_rumdetect2017", "-RD17")
        for label, prefix in specs:
            output.append(
                [
                    dataset_label,
                    label,
                    f"{fmt(row[f'{prefix}_mape_mean'])}±{fmt(row[f'{prefix}_mape_std'])}",
                    fmt(row[f"{prefix}_mae_mean"]),
                    fmt(row[f"{prefix}_rmse_mean"]),
                    fmt(row[f"{prefix}_r2_mean"]),
                ]
            )
    return output


def build_chapter_content(doc: Document) -> list[Any]:
    start_index = body_insert_index(doc)

    add_heading(doc, "5.1实验设置", "Heading 3")
    add_paragraph(
        doc,
        "为验证本文所提出的HeteroRumorDyn模型在谣言早期传播建模、传播动能因子分解以及文本语义贡献分析中的有效性，本章基于统一预处理后的真实传播数据集开展实验。实验任务设置为早期传播规模预测：给定源帖文本、观测窗口内的传播拓扑与时序日志，预测级联最终规模。该任务能够直接检验模型对传播动力学演化趋势的拟合能力，也是后续破圈预警与干预控制模块的基础。"
    )

    add_heading(doc, "5.1.1实验数据", "Heading 4")
    add_paragraph(
        doc,
        "本文主要采用PHEME以及RumDetect2017版本的Twitter15/16数据集进行实验。PHEME数据包含源帖、事件回复、传播结构、相对时间延迟及部分用户画像信息，适用于多模态传播规模预测与传播动能因子分解。RumDetect2017中的Twitter15/16数据包含源推文文本、传播树和节点相对延迟时间，适合用于验证模型在跨数据集场景下的可迁移性及文本编码器贡献。需要说明的是，RumDetect2017公开包不包含非源推文正文、绝对时间戳和完整用户画像，因此本文在该数据集上的文本模态仅使用源推文文本，避免引入未来信息泄漏。"
    )
    add_table(
        doc,
        "表5-1 实验数据集统计信息",
        ["数据集", "样本数", "传播边数", "含源文本样本", "平均级联规模", "可用模态"],
        build_dataset_rows(),
        [1.05, 0.75, 0.9, 0.95, 0.9, 2.15],
    )

    add_heading(doc, "5.1.2基线方法", "Heading 4")
    add_paragraph(
        doc,
        "为全面评估模型性能，本文设置了三类对比方法。第一类为结构统计基线，直接利用级联规模、深度、平均分支因子等统计特征进行分类或回归预测；第二类为传播图基线，包括基于传播结构传播平滑特征的图模型，用于衡量拓扑信息本身的贡献；第三类为HeteroRumorDyn系列模型，包括使用稳定hash文本特征的V1模型、使用MiniLM预训练文本编码器的V1模型、去除文本模态的V1消融模型，以及在PHEME上进一步加入VAE传播动能因子分解与目标匹配文本替换约束的V2/C1模型。"
    )

    add_heading(doc, "5.1.3评价指标", "Heading 4")
    add_paragraph(
        doc,
        "传播规模预测任务采用平均绝对误差（Mean Absolute Error, MAE）、均方根误差（Root Mean Square Error, RMSE）、平均绝对百分比误差（Mean Absolute Percentage Error, MAPE）以及决定系数R2进行评价。其中MAE和RMSE反映预测规模与真实规模之间的绝对偏差，MAPE衡量相对误差，R2衡量模型对级联规模方差的解释能力。除单次结果外，本文对主要模型采用5个随机种子进行复验，并报告均值与标准差；对于文本消融实验，进一步采用配对比较分析MiniLM文本编码器相对于hash文本特征和去文本设置的边际贡献。"
    )

    add_heading(doc, "5.1.4模型参数设置", "Heading 4")
    add_paragraph(
        doc,
        "所有早期传播规模预测实验均采用180分钟观测窗口。数据划分使用stratified seed42划分，以保证训练集、验证集和测试集中的标签分布相对稳定；在防泄漏压力测试中另保留temporal/proxy划分。神经网络训练采用AdamW优化器，学习率为0.001，权重衰减为0.0001，隐藏维度为64，图编码层数为2，dropout设置为0.3，batch size为64，最大训练轮数为30，并根据验证集MAPE进行早停。V2/C1模型中传播动能潜变量维度K设为4，反事实约束采用目标匹配文本替换策略，以降低随机文本扰动造成的无效对照。"
    )

    add_heading(doc, "5.2实验分析", "Heading 3")
    add_paragraph(
        doc,
        "本节从传播规模预测性能、文本语义贡献和模型稳健性三个角度分析实验结果。由于当前实验主要围绕传播动能辨识与早期规模预测展开，因此本章重点验证C1相关模块；破圈预警与闭环控制相关实验将在后续完整系统实验中进一步扩展。"
    )

    add_heading(doc, "5.2.1传播动能因子分解有效性分析", "Heading 4")
    add_paragraph(
        doc,
        "表5-2给出了PHEME数据集上180分钟窗口下的传播规模预测结果。MiniLM增强的V1模型作为强多模态基线，在五个随机种子下取得0.1247±0.0019的MAPE。验证集选择的V2/C1 VAE因子模型在seed42上将测试MAPE进一步降低到0.1231，相对V1均值下降1.30%，说明多模态表征中确实存在可压缩的低维传播动能信号。进一步地，采用内容因子与动力因子显式分离的V2/C1 disentangled模型在五个随机种子下取得0.1243±0.0015的MAPE，略优于V1多种子均值，同时提供了更清晰的因子解释结构。"
    )
    add_table(
        doc,
        "表5-2 PHEME数据集上传播规模预测结果",
        ["模型", "种子数", "MAPE", "MAE", "RMSE", "R2", "相对V1增益"],
        build_pheme_rows(),
        [1.95, 0.55, 1.05, 0.75, 0.8, 0.65, 0.95],
    )
    add_paragraph(
        doc,
        "需要注意的是，V2/C1 VAE因子模型目前为验证集选择的单种子结果，因此更适合作为seed42下的最优参考，而不是完全复验后的稳定估计。相比之下，V2/C1 disentangled模型虽然预测性能提升幅度较小，但其五随机种子结果更加稳健，并能分别刻画文本内容因素与传播动力因素，更符合本文关于传播动能可解释建模的目标。"
    )

    add_heading(doc, "5.2.2文本编码器公平对照分析", "Heading 4")
    add_paragraph(
        doc,
        "为了进一步分析文本语义在早期传播规模预测中的贡献，本文在RumDetect2017 Twitter15/16数据集上构建了公平对照实验。三种设置分别为：使用稳定hash文本特征的V1模型、使用MiniLM预训练文本特征的V1模型，以及在相同模型结构下关闭文本模态的消融模型。三组实验均采用相同的stratified划分、180分钟观测窗口和5个随机种子，保证结果具有可比性。"
    )
    add_table(
        doc,
        "表5-3 RumDetect2017数据集上的文本公平对照结果",
        ["数据集", "文本设置", "MAPE", "MAE", "RMSE", "R2"],
        build_fairness_rows(),
        [1.25, 1.0, 1.05, 0.85, 0.85, 0.7],
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(FIGURE_PATH), width=Inches(6.1))
    add_caption(doc, "图5-1 RumDetect2017数据集上的文本编码器公平对照结果")
    add_paragraph(
        doc,
        "由表5-3和图5-1可知，MiniLM文本编码器在两个数据集上相较hash文本特征均降低了平均MAPE：Twitter15从0.1852降至0.1834，相对下降0.95%；Twitter16从0.1993降至0.1930，相对下降3.18%。其中Twitter16上的文本收益更加稳定，MiniLM在4/5个随机种子上优于hash设置，且去除文本后MAPE增加0.0059，配对t-test的p值为0.0115，表明源推文语义对该数据集的早期传播规模预测具有较明确贡献。"
    )

    add_heading(doc, "5.2.3消融结果与指标差异分析", "Heading 4")
    add_paragraph(
        doc,
        "Twitter15数据集上的文本贡献呈现一定的指标依赖性。虽然MiniLM相较hash文本特征使平均MAPE小幅降低0.0018，但去除文本后的MAPE反而低于完整MiniLM模型0.0037。然而，从MAE、RMSE和R2看，去除文本模型分别退化到8.2308、18.8021和0.8160，弱于完整MiniLM模型的7.8897、17.7634和0.8356。这说明去除文本后的模型可能更倾向于保守预测，从而在相对误差指标上获得局部优势，但其对绝对级联规模的拟合能力下降。"
    )
    add_paragraph(
        doc,
        "上述现象也说明，在不同传播数据集上，文本语义、传播拓扑和早期时序信号的主导程度并不完全一致。Twitter15中早期传播结构和时序增长可能已经解释了较大比例的最终规模变化，文本语义的边际贡献较弱；而Twitter16中源推文语义与传播规模之间的关联更明显，因此MiniLM文本编码器能够提供更稳定的性能增益。"
    )

    add_heading(doc, "5.2.4实验小结", "Heading 4")
    add_paragraph(
        doc,
        "综合上述实验结果，本文模型在传播规模预测任务上能够有效融合文本、拓扑和时序信号，并在PHEME数据集上通过V2/C1因子分解获得更具解释性的低维传播动能表示。RumDetect2017上的跨数据集实验进一步表明，预训练文本语义编码通常优于简单hash文本表示，但其收益会受到数据集传播结构、文本覆盖范围和指标选择的影响。因此，在后续实验中，有必要继续将V2/C1 disentangled模型推广到Twitter15/16数据集，并进一步补充破圈预警与闭环控制模块的完整验证。"
    )

    return collect_new_elements(doc, start_index)


def main() -> None:
    doc = Document(INPUT_DOCX)
    remove_old_empty_experiment_body(doc)
    elements = build_chapter_content(doc)
    move_before_reference(doc, elements)
    doc.save(OUTPUT_DOCX)
    print(json.dumps({"output": str(OUTPUT_DOCX)}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
